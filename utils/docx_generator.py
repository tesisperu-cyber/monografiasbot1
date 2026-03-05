"""
Generador de documentos Word (.docx) para MonografiasBot1
Usa python-docx — no requiere Node.js
"""

from datetime import date
import re


def generar_docx(secciones: dict, contexto: dict, ruta_salida: str) -> bool:
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # ── Página: US Letter, márgenes 2.5 cm ───────────────────────────────
        for section in doc.sections:
            section.page_width    = Cm(21.59)
            section.page_height   = Cm(27.94)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.5)
            section.top_margin    = Cm(2.5)
            section.bottom_margin = Cm(2.5)

        FONT  = 'Times New Roman'
        anio  = date.today().year

        # ── Helpers ───────────────────────────────────────────────────────────
        def fmt(run, size=12, bold=False, italic=False):
            run.font.name   = FONT
            run.font.size   = Pt(size)
            run.font.bold   = bold
            run.font.italic = italic

        def h1(text):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after  = Pt(6)
            p.paragraph_format.left_indent  = Cm(0)
            run = p.add_run(text.upper())
            fmt(run, size=14, bold=True)

        def h2(text):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.left_indent  = Cm(0)
            run = p.add_run(text)
            fmt(run, size=12, bold=True)

        def h3(text):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Cm(0)
            run = p.add_run(text)
            fmt(run, size=12, bold=True, italic=True)

        def body(text):
            if not text.strip():
                return
            p = doc.add_paragraph()
            p.paragraph_format.left_indent       = Cm(0)
            p.paragraph_format.first_line_indent = Cm(1.27)
            p.paragraph_format.space_after        = Pt(0)
            p.paragraph_format.line_spacing       = Pt(24)
            run = p.add_run(text.strip())
            fmt(run, size=12)

        def ref(text):
            if not text.strip():
                return
            p = doc.add_paragraph()
            p.paragraph_format.left_indent       = Cm(1.27)
            p.paragraph_format.first_line_indent = Cm(-1.27)
            p.paragraph_format.space_after        = Pt(4)
            run = p.add_run(text.strip())
            fmt(run, size=12)

        def parse(texto):
            if not texto:
                return
            for linea in texto.split('\n'):
                t = linea.strip()
                if not t:
                    doc.add_paragraph()
                    continue
                if re.match(r'^CAP[IÍ]TULO\s+\d+', t, re.IGNORECASE):
                    h2(t)
                elif re.match(r'^\d+\.\d+[\s.]', t):
                    h3(t)
                elif t == t.upper() and 4 < len(t) < 80 and re.search(r'[A-ZÁÉÍÓÚÑ]', t):
                    h3(t)
                else:
                    body(t)

        # ── PORTADA ───────────────────────────────────────────────────────────
        doc.add_paragraph()
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('MONOGRAFÍA ACADÉMICA')
        fmt(run, size=16, bold=True)

        doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(contexto.get('tema', '').upper())
        fmt(run, size=14, bold=True)

        doc.add_paragraph()
        doc.add_paragraph()

        for label, key in [('Curso', 'curso'), ('Especialidad', 'especialidad'),
                            ('Nivel', 'nivel'), ('Páginas', 'paginas')]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f'{label}: {contexto.get(key, "")}')
            fmt(run, size=12)

        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(anio))
        fmt(run, size=12)

        doc.add_page_break()

        # ── I. INTRODUCCIÓN ───────────────────────────────────────────────────
        h1('I. Introducción')
        parse(secciones.get('introduccion', ''))
        doc.add_page_break()

        # ── II. DESARROLLO ────────────────────────────────────────────────────
        h1('II. Desarrollo')
        parse(secciones.get('desarrollo', ''))
        doc.add_page_break()

        # ── III. CONCLUSIONES ─────────────────────────────────────────────────
        h1('III. Conclusiones y Recomendaciones')
        parse(secciones.get('conclusiones', ''))
        doc.add_page_break()

        # ── IV. REFERENCIAS ───────────────────────────────────────────────────
        h1('IV. Referencias Bibliográficas')
        p = doc.add_paragraph()
        run = p.add_run('(Formato APA 7ma edición — Español, Inglés y Portugués)')
        fmt(run, size=12, italic=True)

        for linea in secciones.get('referencias', '').split('\n'):
            if linea.strip():
                ref(linea)

        doc.save(ruta_salida)
        return True

    except ImportError:
        print("  ⚠ Instala python-docx: pip install python-docx")
        return False
    except Exception as e:
        print(f"  ⚠ Error generando DOCX: {e}")
        return False
