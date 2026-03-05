"""
Generador de documentos Word (.docx) para MonografiasBot1
Usa python-docx — no requiere Node.js
"""

from datetime import date

def generar_docx(secciones: dict, contexto: dict, ruta_salida: str) -> bool:
    """
    Genera un archivo .docx académico con todas las secciones de la monografía.
    Retorna True si tuvo éxito.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import re

        doc = Document()

        # ── Configurar página (US Letter, márgenes 2.5cm) ─────────────────────
        section = doc.sections[0]
        section.page_width  = int(21.59 * 914400 / 25.4)   # 8.5 pulgadas en EMU
        section.page_height = int(27.94 * 914400 / 25.4)   # 11 pulgadas en EMU
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)

        FONT_NAME = 'Times New Roman'
        FONT_SIZE_BODY = Pt(12)
        FONT_SIZE_H1   = Pt(14)
        FONT_SIZE_H2   = Pt(12)

        anio = date.today().year

        def set_font(run, size=FONT_SIZE_BODY, bold=False, italic=False):
            run.font.name  = FONT_NAME
            run.font.size  = size
            run.font.bold  = bold
            run.font.italic = italic

        def add_heading1(text):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after  = Pt(6)
            run = p.add_run(text.upper())
            set_font(run, size=FONT_SIZE_H1, bold=True)
            return p

        def add_heading2(text):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.left_indent  = Cm(1.27)
            run = p.add_run(text)
            set_font(run, size=FONT_SIZE_H2, bold=True)
            return p

        def add_heading3(text):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(1.27)
            run = p.add_run(text)
            set_font(run, size=FONT_SIZE_BODY, bold=True, italic=True)
            return p

        def add_body(text):
            if not text.strip():
                return
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent  = Cm(1.27)
            p.paragraph_format.space_after  = Pt(4)
            # Doble espacio
            p.paragraph_format.line_spacing = Pt(24)
            run = p.add_run(text.strip())
            set_font(run)
            return p

        def add_reference(text):
            if not text.strip():
                return
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent    = Cm(2.54)
            p.paragraph_format.first_line_indent = Cm(-1.27)
            p.paragraph_format.space_after    = Pt(4)
            run = p.add_run(text.strip())
            set_font(run)
            return p

        def parse_and_add(texto):
            """Parsea el texto y agrega párrafos detectando headings."""
            if not texto:
                return
            for linea in texto.split('\n'):
                t = linea.strip()
                if not t:
                    doc.add_paragraph()
                    continue
                if re.match(r'^CAPÍTULO\s+\d+[:.]/i', t, re.IGNORECASE) or \
                   re.match(r'^CAPÍTULO\s+\d+', t, re.IGNORECASE):
                    add_heading2(t)
                elif re.match(r'^\d+\.\d+[\s.]', t):
                    add_heading3(t)
                elif t == t.upper() and 4 < len(t) < 80 and re.search(r'[A-ZÁÉÍÓÚÑ]', t):
                    add_heading3(t)
                else:
                    add_body(t)

        # ── PORTADA ───────────────────────────────────────────────────────────
        doc.add_paragraph()
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('MONOGRAFÍA ACADÉMICA')
        set_font(run, size=Pt(16), bold=True)

        doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(contexto.get('tema', '').upper())
        set_font(run, size=FONT_SIZE_H1, bold=True)

        doc.add_paragraph()
        doc.add_paragraph()

        for label, key in [('Curso', 'curso'), ('Especialidad', 'especialidad'),
                            ('Nivel', 'nivel'), ('Páginas', 'paginas')]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f'{label}: {contexto.get(key, "")}')
            set_font(run)

        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(anio))
        set_font(run)

        doc.add_page_break()

        # ── I. INTRODUCCIÓN ───────────────────────────────────────────────────
        add_heading1('I. Introducción')
        parse_and_add(secciones.get('introduccion', ''))
        doc.add_page_break()

        # ── II. DESARROLLO ────────────────────────────────────────────────────
        add_heading1('II. Desarrollo')
        parse_and_add(secciones.get('desarrollo', ''))
        doc.add_page_break()

        # ── III. CONCLUSIONES ─────────────────────────────────────────────────
        add_heading1('III. Conclusiones y Recomendaciones')
        parse_and_add(secciones.get('conclusiones', ''))
        doc.add_page_break()

        # ── IV. REFERENCIAS ───────────────────────────────────────────────────
        add_heading1('IV. Referencias Bibliográficas')
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.27)
        run = p.add_run('(Formato APA 7ma edición — Español, Inglés y Portugués)')
        set_font(run, italic=True)

        for linea in secciones.get('referencias', '').split('\n'):
            if linea.strip():
                add_reference(linea)
            else:
                doc.add_paragraph()

        # ── Guardar ───────────────────────────────────────────────────────────
        doc.save(ruta_salida)
        return True

    except ImportError:
        print("  ⚠ Instala python-docx: pip install python-docx")
        return False
    except Exception as e:
        print(f"  ⚠ Error generando DOCX: {e}")
        return False
